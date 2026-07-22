"""Render TailoredResume to LaTeX, PDF, plain text, and optional DOCX."""

from __future__ import annotations

import re
import shutil
import subprocess
from datetime import date
from pathlib import Path
from typing import Any

from jinja2 import Environment, FileSystemLoader, select_autoescape

from src.schemas import TailoredResume
from src.settings_loader import load_settings, project_root, resolve_path

_LATEX_SPECIAL = re.compile(r"([\\&%$#_{}~^])")
# Emoji and symbols that break Tectonic / pdfLaTeX on Windows.
_EMOJI_RE = re.compile(
    "["
    "\U0001F300-\U0001FAFF"  # misc symbols & pictographs, emoticons, etc.
    "\U00002700-\U000027BF"  # dingbats
    "\U00002600-\U000026FF"  # misc symbols
    "\U0000FE00-\U0000FE0F"  # variation selectors
    "\U0000200D"             # zero-width joiner
    "\U0001F1E0-\U0001F1FF"  # flags
    "]+",
    flags=re.UNICODE,
)
_PROBLEMATIC_SYMBOLS_RE = re.compile(r"[\U0001F517★☆●○■□◆◇✓✔✗✘♦]")
# Non-ASCII letters/digits outside common Latin-1 punctuation — transliterate or drop.
_NON_LATIN_RE = re.compile(r"[^\x00-\xFF]")

# Normalise common Unicode punctuation to ASCII BEFORE dropping non-Latin-1 chars
# (previously smart quotes/dashes were silently turned into spaces).
_UNICODE_PUNCT = {
    "\u2013": "-", "\u2014": "--",
    "\u2018": "'", "\u2019": "'", "\u201a": "'", "\u201b": "'",
    "\u201c": '"', "\u201d": '"', "\u201e": '"',
    "\u2022": "-", "\u00b7": "-", "\u2027": "-", "\u2043": "-",
    "\u2192": "->", "\u2190": "<-", "\u2794": "->",
    "\u00a0": " ", "\u2009": " ", "\u200a": " ", "\u202f": " ",
    "\u2011": "-", "\u2026": "...", "\u00ad": "",
    "\u2122": "(TM)", "\u00ae": "(R)", "\u00a9": "(C)",
    "\ufb01": "fi", "\ufb02": "fl",
}


def _normalize_unicode(text: str) -> str:
    for src, dst in _UNICODE_PUNCT.items():
        text = text.replace(src, dst)
    return text


def sanitize_for_latex(text: str) -> str:
    """Strip emojis, normalise punctuation, and drop remaining non-Latin-1 chars."""
    if not text:
        return ""
    text = _EMOJI_RE.sub("", text)
    text = _PROBLEMATIC_SYMBOLS_RE.sub("", text)
    text = _normalize_unicode(text)
    text = _NON_LATIN_RE.sub(" ", text)
    return re.sub(r"\s+", " ", text).strip()


def escape_latex(text: str) -> str:
    """Escape LaTeX special characters in plain text."""
    if not text:
        return ""
    text = sanitize_for_latex(text)

    def _replace(match: re.Match[str]) -> str:
        ch = match.group(1)
        mapping = {
            "\\": r"\textbackslash{}",
            "&": r"\&",
            "%": r"\%",
            "$": r"\$",
            "#": r"\#",
            "_": r"\_",
            "{": r"\{",
            "}": r"\}",
            "~": r"\textasciitilde{}",
            "^": r"\textasciicircum{}",
        }
        return mapping[ch]

    return _LATEX_SPECIAL.sub(_replace, text)


def escape_latex_url(text: str) -> str:
    """Escape the characters that break a URL inside \\href{...}."""
    if not text:
        return ""
    text = sanitize_for_latex(text)
    return text.replace("\\", "/").replace("%", r"\%").replace("#", r"\#")


def _jinja_env(templates_dir: Path) -> Environment:
    env = Environment(
        loader=FileSystemLoader(str(templates_dir)),
        autoescape=select_autoescape(default=False),
        trim_blocks=True,
        lstrip_blocks=True,
    )
    env.filters["latex"] = escape_latex
    env.filters["latex_url"] = escape_latex_url
    return env


def tectonic_available() -> bool:
    return shutil.which("tectonic") is not None


def compile_tex_to_pdf(tex_path: Path, out_dir: Path | None = None) -> tuple[Path | None, str | None]:
    """Compile a .tex file to PDF via Tectonic. Returns (pdf_path, error_message)."""
    if not tectonic_available():
        return None, (
            "Tectonic is not installed or not on PATH. "
            "Install it for PDF output (Windows: winget install tectonic-typesetting.tectonic "
            "or scoop install tectonic). .tex and .txt were still written."
        )

    work_dir = out_dir or tex_path.parent
    work_dir.mkdir(parents=True, exist_ok=True)
    try:
        result = subprocess.run(
            ["tectonic", str(tex_path)],
            cwd=str(work_dir),
            capture_output=True,
            text=True,
            timeout=120,
            check=False,
        )
    except (OSError, subprocess.TimeoutExpired) as exc:
        return None, f"Tectonic compilation failed: {exc}"

    pdf_path = tex_path.with_suffix(".pdf")
    if result.returncode != 0 or not pdf_path.exists():
        stderr = (result.stderr or result.stdout or "").strip()
        return None, f"Tectonic compilation failed: {stderr or 'unknown error'}"

    return pdf_path, None


def weasyprint_available() -> bool:
    try:
        import weasyprint  # noqa: F401
    except Exception:
        return False
    return True


def fpdf_available() -> bool:
    try:
        import fpdf  # noqa: F401
    except Exception:
        return False
    return True


def _tailored_to_template_context(tailored: TailoredResume) -> dict[str, Any]:
    return {
        "identity": tailored.identity,
        "summary": tailored.summary,
        "skills": tailored.skills,
        "experience": tailored.experience,
        "projects": tailored.projects,
        "education": tailored.education,
        "certifications": tailored.certifications,
    }


def _html_jinja_env(templates_dir: Path) -> Environment:
    return Environment(
        loader=FileSystemLoader(str(templates_dir)),
        autoescape=select_autoescape(default=True, default_for_string=True),
        trim_blocks=True,
        lstrip_blocks=True,
    )


def render_html(tailored: TailoredResume, settings: dict[str, Any] | None = None) -> str:
    """Render the resume to a standalone HTML string (autoescaped)."""
    cfg = settings or load_settings()
    templates_dir = resolve_path("templates_dir", cfg)
    env = _html_jinja_env(templates_dir)
    template = env.get_template("resume.html.j2")
    return template.render(**_tailored_to_template_context(tailored))


def _html_to_pdf_weasyprint(html_str: str, out_path: Path) -> None:
    from weasyprint import HTML

    out_path.parent.mkdir(parents=True, exist_ok=True)
    HTML(string=html_str).write_pdf(str(out_path))


def render_pdf_fpdf(tailored: TailoredResume, out_path: Path) -> None:
    """Pure-Python PDF (no system deps). Last-resort engine that works anywhere."""
    from fpdf import FPDF
    from fpdf.enums import XPos, YPos

    def clean(text: str) -> str:
        return sanitize_for_latex(text or "")

    left_w = 118.0  # mm; remaining column is right-aligned

    pdf = FPDF(format="Letter", unit="mm")
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.set_margins(16, 14, 16)
    pdf.add_page()

    ident = tailored.identity
    pdf.set_font("Helvetica", "B", 18)
    pdf.cell(0, 8, clean(ident.name or "Resume"), align="C", new_x=XPos.LMARGIN, new_y=YPos.NEXT)

    contact = [
        p for p in (ident.email, ident.phone, ident.location, ident.linkedin, ident.github)
        if p and p != "TODO"
    ]
    if contact:
        pdf.set_font("Helvetica", "", 9)
        pdf.cell(0, 5, clean("   |   ".join(contact)), align="C", new_x=XPos.LMARGIN, new_y=YPos.NEXT)
    pdf.ln(2)

    def section(title: str) -> None:
        pdf.set_font("Helvetica", "B", 11.5)
        pdf.cell(0, 6, clean(title).upper(), new_x=XPos.LMARGIN, new_y=YPos.NEXT)
        y = pdf.get_y()
        pdf.set_draw_color(90, 90, 90)
        pdf.line(pdf.l_margin, y, pdf.w - pdf.r_margin, y)
        pdf.ln(1.5)

    def bullets(items: list[str]) -> None:
        pdf.set_font("Helvetica", "", 10)
        for item in items:
            pdf.set_x(pdf.l_margin + 3)
            pdf.multi_cell(0, 4.7, "-  " + clean(item), new_x=XPos.LMARGIN, new_y=YPos.NEXT)

    if tailored.summary:
        section("Summary")
        pdf.set_font("Helvetica", "", 10)
        pdf.multi_cell(0, 4.8, clean(tailored.summary), new_x=XPos.LMARGIN, new_y=YPos.NEXT)
        pdf.ln(1)

    if tailored.skills:
        section("Skills")
        for category, items in tailored.skills.items():
            pdf.set_font("Helvetica", "B", 10)
            pdf.write(4.8, clean(category) + ": ")
            pdf.set_font("Helvetica", "", 10)
            pdf.multi_cell(0, 4.8, clean(", ".join(items)), new_x=XPos.LMARGIN, new_y=YPos.NEXT)
        pdf.ln(1)

    if tailored.experience:
        section("Experience")
        for job in tailored.experience:
            pdf.set_font("Helvetica", "B", 10.5)
            pdf.cell(left_w, 5, clean(job.company))
            pdf.set_font("Helvetica", "", 9.5)
            pdf.cell(0, 5, clean(f"{job.start} - {job.end}"), align="R",
                     new_x=XPos.LMARGIN, new_y=YPos.NEXT)
            pdf.set_font("Helvetica", "I", 10)
            pdf.cell(left_w, 5, clean(job.title))
            pdf.set_font("Helvetica", "I", 9.5)
            pdf.cell(0, 5, clean(job.location), align="R",
                     new_x=XPos.LMARGIN, new_y=YPos.NEXT)
            bullets(job.bullets)
            pdf.ln(1)

    if tailored.projects:
        section("Projects")
        for project in tailored.projects:
            pdf.set_font("Helvetica", "B", 10.5)
            pdf.cell(0, 5, clean(project.name), new_x=XPos.LMARGIN, new_y=YPos.NEXT)
            bullets(project.bullets)
            pdf.ln(1)

    if tailored.education:
        section("Education")
        for edu in tailored.education:
            detail = clean(edu.degree)
            if edu.field:
                detail += f" | {clean(edu.field)}"
            if edu.gpa:
                detail += f" | CGPA: {clean(edu.gpa)}"
            pdf.set_font("Helvetica", "B", 10.5)
            pdf.cell(left_w, 5, clean(edu.institution))
            pdf.set_font("Helvetica", "", 9.5)
            pdf.cell(0, 5, clean(f"{edu.start} - {edu.end}"), align="R",
                     new_x=XPos.LMARGIN, new_y=YPos.NEXT)
            pdf.set_font("Helvetica", "", 10)
            pdf.multi_cell(0, 4.8, detail, new_x=XPos.LMARGIN, new_y=YPos.NEXT)
            pdf.ln(0.5)

    if tailored.certifications:
        section("Certifications")
        bullets(list(tailored.certifications))

    out_path.parent.mkdir(parents=True, exist_ok=True)
    pdf.output(str(out_path))


def build_pdf(
    tex_path: Path,
    html_str: str,
    tailored: TailoredResume,
    pdf_path: Path,
    settings: dict[str, Any] | None = None,
) -> tuple[Path | None, str | None, str]:
    """Produce a PDF using the best available engine.

    Order (unless overridden by ``features.pdf_engine``): Tectonic -> WeasyPrint -> fpdf2.
    Returns ``(pdf_path_or_None, note_or_None, engine_name)``.
    """
    cfg = settings or load_settings()
    engine = str(cfg.get("features", {}).get("pdf_engine", "auto")).lower()
    errors: list[str] = []

    if engine in ("auto", "latex", "tectonic"):
        if tectonic_available():
            compiled, err = compile_tex_to_pdf(tex_path, pdf_path.parent)
            if compiled:
                return compiled, None, "tectonic"
            if err:
                errors.append(err)
        elif engine != "auto":
            errors.append("Tectonic requested but not installed.")

    if engine in ("auto", "html", "weasyprint"):
        if weasyprint_available():
            try:
                _html_to_pdf_weasyprint(html_str, pdf_path)
                return pdf_path, None, "weasyprint"
            except Exception as exc:  # noqa: BLE001
                errors.append(f"WeasyPrint failed: {exc}")
        elif engine != "auto":
            errors.append("WeasyPrint requested but not installed.")

    if engine in ("auto", "fpdf", "fpdf2"):
        if fpdf_available():
            try:
                render_pdf_fpdf(tailored, pdf_path)
                return pdf_path, None, "fpdf2"
            except Exception as exc:  # noqa: BLE001
                errors.append(f"fpdf2 failed: {exc}")
        elif engine != "auto":
            errors.append("fpdf2 requested but not installed.")

    note = (
        "No PDF engine succeeded. Install one of: Tectonic (best, LaTeX; Windows: "
        "winget install tectonic-typesetting.tectonic), WeasyPrint, or fpdf2. "
        ".tex, .html and .txt were still written."
    )
    if errors:
        note += "\nAttempts:\n- " + "\n- ".join(errors)
    return None, note, ""


def resume_to_plain_text(tailored: TailoredResume) -> str:
    """ATS-friendly plain-text export from a TailoredResume."""
    # Clean bullets that may contain emoji (e.g. from parsed PDF).
    def _clean(s: str) -> str:
        return sanitize_for_latex(s) if s else s

    lines: list[str] = []
    ident = tailored.identity

    if ident.name:
        lines.append(ident.name.upper())
    contact: list[str] = []
    for part in (ident.email, ident.phone, ident.location, ident.linkedin, ident.github):
        if part and part != "TODO":
            contact.append(part)
    if contact:
        lines.append(" | ".join(contact))
    lines.append("")

    if tailored.summary:
        lines.append("SUMMARY")
        lines.append(_clean(tailored.summary))
        lines.append("")

    if tailored.skills:
        lines.append("SKILLS")
        for category, items in tailored.skills.items():
            lines.append(f"{category}: {', '.join(items)}")
        lines.append("")

    if tailored.experience:
        lines.append("EXPERIENCE")
        for job in tailored.experience:
            lines.append(f"{job.title} | {job.company} | {job.location}")
            lines.append(f"{job.start} - {job.end}")
            for bullet in job.bullets:
                lines.append(f"- {_clean(bullet)}")
            lines.append("")

    if tailored.projects:
        lines.append("PROJECTS")
        for project in tailored.projects:
            tech = f" ({', '.join(project.tech)})" if project.tech else ""
            lines.append(f"{project.name}{tech}")
            for bullet in project.bullets:
                lines.append(f"- {_clean(bullet)}")
            lines.append("")

    if tailored.education:
        lines.append("EDUCATION")
        for edu in tailored.education:
            detail = edu.degree
            if edu.field:
                detail = f"{detail} | {edu.field}"
            if edu.gpa:
                detail = f"{detail} | CGPA: {edu.gpa}"
            lines.append(f"{edu.institution} | {edu.location}")
            lines.append(f"{detail} | {edu.start} - {edu.end}")
            lines.append("")

    if tailored.certifications:
        lines.append("CERTIFICATIONS")
        for cert in tailored.certifications:
            lines.append(f"- {cert}")

    return "\n".join(lines).strip() + "\n"


def _render_docx(tailored: TailoredResume, out_path: Path) -> None:
    from docx import Document
    from docx.shared import Pt

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    ident = tailored.identity
    heading = doc.add_heading(ident.name or "Resume", level=0)
    heading.alignment = 1

    contact_bits = [
        p
        for p in (ident.email, ident.phone, ident.location, ident.linkedin, ident.github)
        if p and p != "TODO"
    ]
    if contact_bits:
        doc.add_paragraph(" | ".join(contact_bits))

    def add_section(title: str) -> None:
        doc.add_heading(title, level=1)

    if tailored.summary:
        add_section("Summary")
        doc.add_paragraph(tailored.summary)

    if tailored.skills:
        add_section("Skills")
        for category, items in tailored.skills.items():
            doc.add_paragraph(f"{category}: {', '.join(items)}")

    if tailored.experience:
        add_section("Experience")
        for job in tailored.experience:
            doc.add_paragraph(
                f"{job.title} — {job.company}, {job.location} ({job.start}–{job.end})",
                style="List Bullet",
            )
            for bullet in job.bullets:
                doc.add_paragraph(bullet, style="List Bullet 2")

    if tailored.projects:
        add_section("Projects")
        for project in tailored.projects:
            doc.add_paragraph(project.name, style="List Bullet")
            for bullet in project.bullets:
                doc.add_paragraph(bullet, style="List Bullet 2")

    if tailored.education:
        add_section("Education")
        for edu in tailored.education:
            doc.add_paragraph(
                f"{edu.institution} — {edu.degree} ({edu.start}–{edu.end})",
                style="List Bullet",
            )

    if tailored.certifications:
        add_section("Certifications")
        for cert in tailored.certifications:
            doc.add_paragraph(cert, style="List Bullet")

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))


def render_resume(
    tailored: TailoredResume,
    out_dir: Path | str,
    basename: str,
    formats: list[str] | None = None,
    settings: dict[str, Any] | None = None,
) -> dict[str, Path]:
    """Render tailored resume to .tex, .txt, optional PDF/DOCX.

    Returns a dict of format -> Path. On missing Tectonic, ``pdf`` is omitted and
    ``pdf_note`` contains an actionable install message.
    """
    cfg = settings or load_settings()
    out_path = Path(out_dir)
    out_path.mkdir(parents=True, exist_ok=True)

    export_cfg = cfg.get("export", {})
    default_formats = list(export_cfg.get("formats", ["pdf", "txt"]))
    requested = formats or default_formats
    docx_enabled = bool(export_cfg.get("docx", False))

    templates_dir = resolve_path("templates_dir", cfg)
    env = _jinja_env(templates_dir)
    template = env.get_template("resume.tex.j2")
    context = _tailored_to_template_context(tailored)

    tex_path = out_path / f"{basename}.tex"
    tex_path.write_text(template.render(**context), encoding="utf-8")
    outputs: dict[str, Path] = {"tex": tex_path}

    txt_path = out_path / f"{basename}.txt"
    txt_path.write_text(resume_to_plain_text(tailored), encoding="utf-8")
    outputs["txt"] = txt_path

    html_str = render_html(tailored, cfg)
    html_path = out_path / f"{basename}.html"
    html_path.write_text(html_str, encoding="utf-8")
    outputs["html"] = html_path

    if "pdf" in requested:
        pdf_target = out_path / f"{basename}.pdf"
        pdf_path, pdf_error, _engine = build_pdf(tex_path, html_str, tailored, pdf_target, cfg)
        if pdf_path:
            outputs["pdf"] = pdf_path
        elif pdf_error:
            note_path = out_path / f"{basename}.pdf_note.txt"
            note_path.write_text(pdf_error + "\n", encoding="utf-8")
            outputs["pdf_note"] = note_path

    if docx_enabled and ("docx" in requested or docx_enabled):
        docx_path = out_path / f"{basename}.docx"
        _render_docx(tailored, docx_path)
        outputs["docx"] = docx_path

    return outputs


def profile_to_tailored_sample(profile_dict: dict[str, Any], job_title: str, company: str = "") -> TailoredResume:
    """Helper: wrap a MasterProfile dict as TailoredResume with mock meta (for tests)."""
    from src.schemas import MasterProfile, TailoredResumeMeta

    profile = MasterProfile.from_dict(profile_dict)
    meta = TailoredResumeMeta(
        job_title=job_title,
        company=company or "Example Corp",
        jd_keywords=["UVM", "SystemVerilog", "verification"],
        matched_keywords=["SystemVerilog"],
        missing_keywords=["UVM"],
        model_used="test",
        generated_at=date.today().isoformat(),
    )
    return TailoredResume(
        identity=profile.identity,
        summary=profile.summary,
        skills=profile.skills,
        experience=list(profile.experience),
        projects=list(profile.projects),
        education=list(profile.education),
        certifications=list(profile.certifications),
        achievements=list(profile.achievements),
        bullet_bank=list(profile.bullet_bank),
        meta=meta,
    )
